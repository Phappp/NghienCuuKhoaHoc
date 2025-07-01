import sys
import json
from docx import Document

def extract_text(docx_path):
    try:
        doc = Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        # Đọc thêm text trong bảng
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += "\n" + cell.text
        return {"text": text, "confidence": 1.0, "error": None}
    except Exception as e:
        return {"text": None, "confidence": 0, "error": str(e)}

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"text": None, "confidence": 0, "error": "No file path provided"}))
        sys.exit(1)
    result = extract_text(sys.argv[1])
    print(json.dumps(result)) 